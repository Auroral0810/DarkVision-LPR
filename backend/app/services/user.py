from sqlalchemy import func
from sqlalchemy.orm import Session, joinedload, aliased
from app.models.user import User, UserProfile, UserMembership, RealNameVerification
from app.models.role import AdminRole, Role
from app.schemas.user import UserDetailInfo as UserDetail
from app.schemas.admin.user import AdminSelfInfo, UserListParams, AdminUserCreate, AdminUserUpdate
from typing import Optional, List, Tuple
from app.core.security import get_password_hash
from datetime import datetime, date
import secrets
import string
import logging
import csv
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_user_by_id(db: Session, user_id: int) -> Optional[User]:
    """
    通过ID获取用户模型实例
    """
    return db.query(User).filter(User.id == user_id).first()

def get_user_detail_info(db: Session, user_id: int) -> Optional[UserDetail]:
    """
    获取用户详细信息（包含角色、会员、认证等信息）
    """
    user = db.query(User).options(
        joinedload(User.admin_roles).joinedload(AdminRole.role)
    ).filter(User.id == user_id).first()

    if not user:
        return None

    # 构建基础信息
    user_info = {
        "id": user.id,
        "phone": user.phone,
        "nickname": user.nickname,
        "email": user.email,
        "avatar_url": user.avatar_url,
        "user_type": user.user_type.value if hasattr(user.user_type, 'value') else user.user_type,
        "status": user.status.value if hasattr(user.status, 'value') else user.status,
        "created_at": user.created_at,
        "last_login_at": user.last_login_at,
        "last_login_ip": user.last_login_ip,
        "banned_reason": user.banned_reason,
        "banned_until": user.banned_until,
        "roles": []  # 默认为空列表
    }

    # 获取上级账号信息
    if user.parent_id:
        parent = db.query(User).filter(User.id == user.parent_id).first()
        if parent:
            user_info["parent_id"] = parent.id
            user_info["parent_nickname"] = parent.nickname

    # 计算统计数据
    from app.models.recognition import RecognitionRecord
    total_rec = db.query(func.count(RecognitionRecord.id)).filter(RecognitionRecord.user_id == user_id).scalar()
    user_info["total_recognition_count"] = total_rec or 0
    
    last_rec = db.query(RecognitionRecord.created_at).filter(RecognitionRecord.user_id == user_id).order_by(RecognitionRecord.created_at.desc()).first()
    if last_rec:
        user_info["last_recognition_at"] = last_rec[0]

    user_info["is_deleted"] = user.deleted_at is not None

    # 获取角色信息 (如果是管理员)
    if user.user_type == 'admin':
        # 调试输出
        print(f"DEBUG: Checking roles for user {user_id}")
        import logging
        logger = logging.getLogger("darkvision")
        
        # 通过 AdminRole 关联查询角色名称
        # user.admin_roles 是 relationship 列表，每个元素是 AdminRole 对象，包含 role 对象
        roles = []
        if user.admin_roles:
            print(f"DEBUG: Found {len(user.admin_roles)} admin_roles entries")
            for ar in user.admin_roles:
                print(f"DEBUG: Processing AdminRole: {ar}, Role: {ar.role}")
                if ar.role:
                    roles.append(ar.role.name)
                else:
                    logger.warning(f"AdminRole found but no Role associated for user {user_id}")
        else:
             print("DEBUG: No admin_roles found (empty list)")

        # 如果没有分配具体角色但类型是admin，赋予默认ROOT角色作为容错（可选）
        if not roles:
            print("DEBUG: No specific roles found, defaulting to ROOT")
            roles = ["ROOT"]
            
        user_info["roles"] = roles
    else:
        # 普通用户无特定管理角色
        user_info["roles"] = []

    # 获取会员信息
    if user.membership:
        user_info.update({
            "membership_type": user.membership.membership_type,
            "membership_expire_date": user.membership.expire_date,
            "is_membership_active": bool(user.membership.is_active)
        })

    # 获取个人资料
    if user.profile:
        birthday_str = None
        if user.profile.birthday:
            if isinstance(user.profile.birthday, (datetime, date)):
                birthday_str = user.profile.birthday.strftime('%Y-%m-%d')
            else:
                birthday_str = str(user.profile.birthday)
                
        user_info.update({
            "gender": user.profile.gender,
            "birthday": birthday_str,
            "address": user.profile.address,
            "real_name": user.profile.real_name
        })
        
    # 获取实名认证状态
    if user.real_name_verification:
        user_info.update({
            "is_verified": user.real_name_verification.status == 'approved',
            "verification_status": user.real_name_verification.status,
            "reject_reason": user.real_name_verification.reject_reason,
            "id_card_front": user.real_name_verification.id_card_front,
            "id_card_back": user.real_name_verification.id_card_back,
            "face_photo": user.real_name_verification.face_photo
        })
    else:
         user_info.update({
            "is_verified": False,
            "verification_status": None
        })

    # 计算配额信息 (从 recognition_statistics 表查询今日已用)
    from app.models.recognition import RecognitionStatistic
    today = date.today()
    stat = db.query(RecognitionStatistic).filter(
        RecognitionStatistic.user_id == user_id,
        RecognitionStatistic.stat_date == today
    ).first()
    
    used_today = stat.daily_count if stat else 0
    
    # 根据用户类型设置限额 (通常应从配置表读取，这里根据类型给默认值)
    quota_map = {
        'free': 100,
        'vip': 2000,
        'enterprise': 100000,
        'admin': 999999
    }
    utype = user.user_type.value if hasattr(user.user_type, 'value') else user.user_type
    daily_quota = quota_map.get(utype, 100)
    
    # 企业主账号逻辑
    is_main = (user.user_type == 'enterprise' and user.parent_id is None)
    sub_count = 0
    if is_main:
        sub_count = db.query(func.count(User.id)).filter(User.parent_id == user.id).scalar() or 0
    
    user_info.update({
        "daily_quota": daily_quota,
        "used_quota_today": used_today,
        "remaining_quota_today": max(0, daily_quota - used_today),
        "is_enterprise_main": is_main,
        "sub_account_count": sub_count
    })

    return UserDetail(**user_info)

def get_admin_detail_info(db: Session, user_id: int) -> Optional[AdminSelfInfo]:
    """
    获取管理员详细信息（包含个人资料、认证状态及管理角色信息）
    """
    # 复用 get_user_detail_info 获取完整数据
    detail = get_user_detail_info(db, user_id)
    if not detail:
        return None
        
    # 将 UserDetailInfo 映射到 AdminSelfInfo (会自动忽略不匹配的字段如配额、会员详情等)
    return AdminSelfInfo(**detail.model_dump())

def list_users_for_admin(
    db: Session, 
    params: UserListParams
) -> Tuple[List[dict], int]:
    """
    管理员分页查询用户列表 (增强版)
    """
    ParentUser = aliased(User)
    
    # 基础查询，关联主账号以获取昵称，并也预加载管理员角色
    query = db.query(User, ParentUser.nickname.label("parent_nickname")).\
        outerjoin(ParentUser, User.parent_id == ParentUser.id).\
        options(joinedload(User.admin_roles).joinedload(AdminRole.role))
    
    # 默认排除管理员账号
    query = query.filter(User.user_type != 'admin')
    
    if params.keyword:
        keyword = f"%{params.keyword}%"
        query = query.filter(
            (User.nickname.ilike(keyword)) | 
            (User.phone.ilike(keyword)) | 
            (User.email.ilike(keyword))
        )
    
    if params.user_type:
        query = query.filter(User.user_type == params.user_type)
        
    if params.status:
        query = query.filter(User.status == params.status)
        
    if params.start_date:
        try:
            start_dt = datetime.strptime(params.start_date, '%Y-%m-%d')
            query = query.filter(User.created_at >= start_dt)
        except ValueError:
            pass
            
    if params.end_date:
        try:
            end_dt = datetime.strptime(params.end_date, '%Y-%m-%d').replace(hour=23, minute=59, second=59)
            query = query.filter(User.created_at <= end_dt)
        except ValueError:
            pass
            
    total = query.with_entities(func.count(User.id)).scalar()
    results = query.order_by(User.created_at.desc())\
                 .offset((params.page - 1) * params.page_size)\
                 .limit(params.page_size).all()
    
    # 批量获取识别次数统计，避免 N+1 查询
    from app.models.recognition import RecognitionRecord
    user_ids = [u[0].id for u in results]
    rec_counts = {}
    if user_ids:
        rec_stats = db.query(
            RecognitionRecord.user_id, 
            func.count(RecognitionRecord.id).label("count")
        ).filter(RecognitionRecord.user_id.in_(user_ids)).group_by(RecognitionRecord.user_id).all()
        rec_counts = {rs.user_id: rs.count for rs in rec_stats}

    enriched_users = []
    for user, parent_nickname in results:
        roles = []
        for ar in user.admin_roles:
            if ar.role:
                roles.append({
                    "id": ar.role.id,
                    "name": ar.role.name,
                    "display_name": ar.role.description # 假设这里存的是显示名
                })
        
        # 转换并补全字段
        user_data = {
            "id": user.id,
            "phone": user.phone,
            "nickname": user.nickname,
            "email": user.email,
            "avatar_url": user.avatar_url,
            "user_type": user.user_type.value if hasattr(user.user_type, 'value') else user.user_type,
            "status": user.status.value if hasattr(user.status, 'value') else user.status,
            "last_login_at": user.last_login_at,
            "last_login_ip": user.last_login_ip,
            "created_at": user.created_at,
            "banned_reason": user.banned_reason,
            "banned_until": user.banned_until,
            "roles": roles,
            "is_admin": user.user_type == 'admin',
            "parent_id": user.parent_id,
            "parent_nickname": parent_nickname,
            "is_online": False, # 默认False，下面补全
            "total_recognition_count": rec_counts.get(user.id, 0),
            "total_order_amount": 0.0     # TODO: 后续关联订单表
        }
        
        # 补全在线状态
        from app.services.online_user_service import is_user_online
        user_data["is_online"] = is_user_online(user.id)
        enriched_users.append(user_data)
                 
    return enriched_users, total

def admin_create_user(db: Session, user_in: AdminUserCreate) -> User:
    """
    管理员创建用户
    """
    db_user = User(
        phone=user_in.phone,
        nickname=user_in.nickname,
        email=user_in.email,
        hashed_password=get_password_hash(user_in.password),
        user_type=user_in.user_type,
        status="active"
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    return db_user

def admin_update_user(db: Session, user_db: User, user_in: AdminUserUpdate) -> User:
    """
    管理员更新用户信息
    """
    update_data = user_in.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(user_db, field, value)
    
    db.add(user_db)
    db.commit()
    db.refresh(user_db)
    return user_db

def batch_delete_users(db: Session, user_ids: List[int]) -> int:
    """
    批量删除用户
    """
    deleted_count = db.query(User).filter(User.id.in_(user_ids)).delete(synchronize_session=False)
    db.commit()
    return deleted_count

def ban_user(db: Session, user_db: User, reason: str, banned_until: Optional[datetime] = None) -> User:
    """
    封禁用户
    """
    user_db.status = "banned"
    user_db.banned_reason = reason
    user_db.banned_until = banned_until
    
    db.add(user_db)
    db.commit()
    db.refresh(user_db)
    return user_db

def unban_user(db: Session, user_db: User) -> User:
    """
    解封用户
    """
    user_db.status = "active"
    user_db.banned_reason = None
    user_db.banned_until = None
    
    db.add(user_db)
    db.commit()
    db.refresh(user_db)
    return user_db

def reset_user_password(db: Session, user_db: User) -> str:
    """
    重置用户密码为随机8位字符串
    """
    alphabet = string.ascii_letters + string.digits
    new_password = ''.join(secrets.choice(alphabet) for i in range(8))
    
    user_db.hashed_password = get_password_hash(new_password)
    db.add(user_db)
    db.commit()
    return new_password

def generate_user_csv(users: List[dict]) -> str:
    """
    生成用户列表 CSV 字符串 (更新字段集)
    """
    output = io.StringIO()
    # 增加 UTF-8 BOM 以支持 Excel 直接打开不乱码
    output.write('\ufeff')
    writer = csv.writer(output)
    
    # 写入表头 (根据用户要求的字段顺序)
    writer.writerow([
        'ID', '手机号', '昵称', '邮箱', '头像地址', '用户类型', '状态', 
        '最后登录时间', '最后登录IP', '注册时间', '是否在线', 
        '封禁原因', '封禁截止', '角色', '是否管理员', 
        '父账号ID', '父账号昵称', '累计识别次数', '累计消费金额'
    ])
    
    # 写入数据
    for u in users:
        # 处理 roles 列表
        roles_str = ', '.join([r['display_name'] for r in u.get('roles', [])]) if isinstance(u.get('roles'), list) else ''
        
        writer.writerow([
            u.get('id'),
            u.get('phone') or '',
            u.get('nickname'),
            u.get('email') or '',
            u.get('avatar_url') or '',
            u.get('user_type'),
            u.get('status'),
            u.get('last_login_at').strftime('%Y-%m-%d %H:%M:%S') if u.get('last_login_at') else '',
            u.get('last_login_ip') or '',
            u.get('created_at').strftime('%Y-%m-%d %H:%M:%S') if u.get('created_at') else '',
            '在线' if u.get('is_online') else '离线',
            u.get('banned_reason') or '',
            u.get('banned_until').strftime('%Y-%m-%d %H:%M:%S') if u.get('banned_until') else '',
            roles_str,
            '是' if u.get('is_admin') else '否',
            u.get('parent_id') or '',
            u.get('parent_nickname') or '',
            u.get('total_recognition_count', 0),
            u.get('total_order_amount', 0)
        ])
    
    content = output.getvalue()
    output.close()
    return content

def generate_user_detail_docx(detail: UserDetail) -> io.BytesIO:
    """
    生成用户详情 Word 文档 (深度美化版)
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    
    # --- 全局字体设置 ---
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = '微软雅黑'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # --- 标题装饰 ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f'DarkVision-LPR 用户档案报告')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80) # 深蓝色

    doc.add_paragraph().add_run(f'USER PROFILE: {detail.nickname.upper()}').bold = True
    doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # 画一条水平线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 50)
    run.font.color.rgb = RGBColor(189, 195, 199)

    def add_section_header(text):
        h = doc.add_paragraph()
        run = h.add_run(f'  {text}')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # 背景颜色 (需要 Oxml 处理)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '2C3E50') # 深蓝色背景
        h._element.get_or_add_pPr().append(shading_elm)

    # --- 1. 基础信息卡片 ---
    add_section_header('账户概览 / ACCOUNT OVERVIEW')
    
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.5)

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value if value is not None else 'N/A')
        # 美化标签
        for p in row[0].paragraphs:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(127, 140, 141)

    add_row('用户 ID', f'#{detail.id}')
    add_row('昵称 (Nickname)', detail.nickname)
    add_row('手机 (Phone)', detail.phone)
    add_row('邮箱 (Email)', detail.email)
    add_row('用户类型', detail.user_type.upper())
    
    status_text = '正常 (NORMAL)' if detail.status == 'active' else '封禁 (BANNED)'
    add_row('当前状态', status_text)
    
    # Determine if the user is an admin
    is_admin = detail.user_type == 'admin' or (hasattr(detail, 'is_admin') and detail.is_admin)

    # --- Conditional Sections based on user type ---
    if is_admin:
        doc.add_paragraph() # 空行
        add_section_header('管理员权限信息 / ADMIN PRIVILEGES')
        
        role_labels = [r for r in detail.roles] if detail.roles else ['无特殊角色']
        p = doc.add_paragraph()
        p.add_run('当前拥有角色权限: ').bold = True
        for role in role_labels:
            p_role = doc.add_paragraph(f'• {role}')
            p_role.paragraph_format.left_indent = Inches(0.5)
            
        doc.add_paragraph()
        doc.add_paragraph('管理配额: 系统管理员拥有极高配额且不受常规限制。').italic = True
    else:
        # --- 2. 状态与配额 (普通用户) ---
        doc.add_paragraph() # 空行
        add_section_header('账户状态与配额 / ACCOUNT STATUS & QUOTA')
        
        p = doc.add_paragraph()
        p.add_run('当前状态: ').bold = True
        status_run = p.add_run(detail.status.upper())
        if detail.status == 'active':
            status_run.font.color.rgb = RGBColor(0, 128, 0) # Green
        elif detail.status == 'banned':
            status_run.font.color.rgb = RGBColor(255, 0, 0) # Red
            
        doc.add_paragraph(f'最后登录时间: {detail.last_login_at.strftime("%Y-%m-%d %H:%M") if detail.last_login_at else "从未登录"}')
        doc.add_paragraph(f'最后登录 IP: {detail.last_login_ip or "未知"}')
        
        doc.add_paragraph('今日识别配额进度:').bold = True
        quota_desc = f'已使用: {detail.used_quota_today} / 总计: {detail.daily_quota} ({detail.remaining_quota_today} 剩余)'
        doc.add_paragraph(quota_desc)

        # --- 3. 业务统计 (仅普通用户显示) ---
        doc.add_paragraph() # 空行
        add_section_header('服务使用统计 / USAGE STATISTICS')
        
        stats_table = doc.add_table(rows=2, cols=2)
        stats_table.style = 'Table Grid' # Changed from 'Light Shading Accent 1' to match existing style
        
        cells = stats_table.rows[0].cells
        cells[0].text = '累计识别数量'
        cells[1].text = f'{detail.total_recognition_count} 次'
        
        cells = stats_table.rows[1].cells
        cells[0].text = '今日已用配额'
        cells[1].text = f'{detail.used_quota_today} / {detail.daily_quota}'
        
        # 字体加粗和颜色
        for row in stats_table.rows:
            for cell in row.cells:
                for p_cell in cell.paragraphs: # Renamed p to p_cell to avoid conflict
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if len(p_cell.runs) > 0:
                        p_cell.runs[0].font.bold = True

    # --- 4. 实名与认证 ---
    doc.add_paragraph() # 空行
    add_section_header('身份识别信息 / IDENTITY INFO')
    
    p = doc.add_paragraph()
    p.add_run('真实姓名: ').bold = True
    p.add_run(detail.real_name or '未通过实名认证')
    
    p = doc.add_paragraph()
    p.add_run('居住地址: ').bold = True
    p.add_run(detail.address or '未填写')

    # --- 5. 管理备注 --- (Original section 4, now 5)
    if detail.status == 'banned':
        doc.add_paragraph()
        add_section_header('系统管控详情 / SYSTEM CONTROL')
        doc.add_paragraph(f'违规原因: {detail.banned_reason or "系统自动封禁"}').runs[0].font.color.rgb = RGBColor(231, 76, 60)
        doc.add_paragraph(f'封禁截止: {detail.banned_until.strftime("%Y-%m-%d %H:%M") if detail.banned_until else "永久"}')

    # 页脚美化
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('\n(EOF) DarkVision-LPR 管理系统机密文档')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(149, 165, 166)

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target
